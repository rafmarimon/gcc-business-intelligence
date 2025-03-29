#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Document Processor Module for Market Intelligence Platform.

This module handles processing of uploaded documents (PDF, DOCX, CSV, XLSX)
and extracts their text content for analysis.
"""

import logging
import os
import re
import pandas as pd
from datetime import datetime
from typing import Dict, Tuple, Any, List, Optional

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def process_document(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Process a document and extract its text content.
    
    Args:
        file_path: Path to the file to process
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return _process_pdf(file_path)
    elif file_extension == '.docx':
        return _process_docx(file_path)
    elif file_extension == '.txt':
        return _process_txt(file_path)
    elif file_extension in ('.csv', '.xlsx', '.xls'):
        return _process_tabular(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def _process_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from a PDF file."""
    try:
        import pypdf
        
        extracted_text = ""
        metadata = {}
        
        # Open the PDF file
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            # Extract metadata
            metadata = {
                'title': pdf_reader.metadata.get('/Title', ''),
                'author': pdf_reader.metadata.get('/Author', ''),
                'creator': pdf_reader.metadata.get('/Creator', ''),
                'producer': pdf_reader.metadata.get('/Producer', ''),
                'subject': pdf_reader.metadata.get('/Subject', ''),
                'pages': len(pdf_reader.pages),
                'creation_date': pdf_reader.metadata.get('/CreationDate', '')
            }
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                extracted_text += page.extract_text() + "\n\n"
        
        # Clean up text
        extracted_text = _clean_text(extracted_text)
        
        logger.info(f"Processed PDF: {os.path.basename(file_path)} - {len(extracted_text)} characters extracted")
        return extracted_text, metadata
        
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        raise

def _process_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from a DOCX file."""
    try:
        import docx
        
        doc = docx.Document(file_path)
        
        # Extract metadata
        metadata = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'created': doc.core_properties.created.isoformat() if doc.core_properties.created else '',
            'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else '',
            'comments': doc.core_properties.comments or '',
            'paragraphs': len(doc.paragraphs)
        }
        
        # Extract text from paragraphs
        extracted_text = ""
        for para in doc.paragraphs:
            extracted_text += para.text + "\n"
        
        # Clean up text
        extracted_text = _clean_text(extracted_text)
        
        logger.info(f"Processed DOCX: {os.path.basename(file_path)} - {len(extracted_text)} characters extracted")
        return extracted_text, metadata
        
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        raise

def _process_txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from a plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            extracted_text = file.read()
        
        # Get file stats
        file_stats = os.stat(file_path)
        
        metadata = {
            'file_size': file_stats.st_size,
            'modification_time': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            'creation_time': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
            'line_count': extracted_text.count('\n') + 1,
            'word_count': len(extracted_text.split())
        }
        
        # Clean up text
        extracted_text = _clean_text(extracted_text)
        
        logger.info(f"Processed TXT: {os.path.basename(file_path)} - {len(extracted_text)} characters extracted")
        return extracted_text, metadata
        
    except Exception as e:
        logger.error(f"Error processing TXT: {str(e)}")
        raise

def _process_tabular(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from a CSV or Excel file."""
    try:
        # Determine the file type
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Read the file with pandas
        if file_extension == '.csv':
            df = pd.read_csv(file_path, encoding='utf-8', errors='replace')
        else:  # Excel files
            df = pd.read_excel(file_path)
        
        # Extract metadata
        metadata = {
            'rows': len(df),
            'columns': len(df.columns),
            'column_names': df.columns.tolist(),
            'file_type': file_extension[1:]  # Remove the dot
        }
        
        # Convert to formatted text
        # First, add the column names
        extracted_text = "# " + " | ".join(str(col) for col in df.columns) + "\n"
        extracted_text += "-" * 80 + "\n"
        
        # Then add the data
        for _, row in df.head(100).iterrows():  # Limit to first 100 rows to avoid huge texts
            extracted_text += " | ".join(str(val) for val in row.values) + "\n"
        
        if len(df) > 100:
            extracted_text += f"\n... (showing 100 of {len(df)} rows) ...\n"
        
        # Add summary statistics for numeric columns
        extracted_text += "\n## Data Summary\n"
        for col in df.columns:
            if pd.api.types.is_numeric_dtype(df[col]):
                extracted_text += f"{col}:\n"
                extracted_text += f"  Min: {df[col].min()}\n"
                extracted_text += f"  Max: {df[col].max()}\n"
                extracted_text += f"  Mean: {df[col].mean()}\n"
                extracted_text += f"  Median: {df[col].median()}\n"
                extracted_text += "\n"
        
        # Clean up text
        extracted_text = _clean_text(extracted_text)
        
        logger.info(f"Processed tabular file: {os.path.basename(file_path)} - {len(extracted_text)} characters extracted")
        return extracted_text, metadata
        
    except Exception as e:
        logger.error(f"Error processing tabular file: {str(e)}")
        raise

def _clean_text(text: str) -> str:
    """Clean up extracted text."""
    # Replace multiple newlines with a single one
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Replace multiple spaces with a single one
    text = re.sub(r' +', ' ', text)
    
    # Remove any non-printable characters
    text = re.sub(r'[^\x20-\x7E\n\t]', '', text)
    
    return text.strip()

if __name__ == "__main__":
    # Simple test
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"Processing: {file_path}")
        
        extracted_text, metadata = process_document(file_path)
        
        print("\n=== METADATA ===")
        for key, value in metadata.items():
            print(f"{key}: {value}")
        
        print("\n=== EXTRACTED TEXT SAMPLE ===")
        print(extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text)
        print(f"\nTotal characters: {len(extracted_text)}")
    else:
        print("Usage: python document_processor.py <file_path>") 